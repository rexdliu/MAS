import streamlit as st
import pandas as pd
import plotly.express as px
from jamaibase import JamAI, protocol as p
import random
from docx import Document
import string
from io import BytesIO


# 初始化 JamAI 对象
jamai = JamAI(api_key="jamai_sk_af32e3adbe4a3f365322576ed8ecc524a5e678e42ae2bc4d", project_id="proj_319b40160d3ccdca5f621f06")

# 生成随机文件名的函数
def generate_random_filename(extension=".docx"):
    random_str = ''.join(random.choices(string.ascii_letters + string.digits, k=10))
    return f"nutrition_report_{random_str}{extension}"

# 设置页面配置
st.set_page_config(page_title="Daily Meal Analysis", page_icon="🍽️")

# 初始化 session state
if 'food_entries' not in st.session_state:
    st.session_state.food_entries = []
if 'user_entries' not in st.session_state:
    st.session_state.user_entries = []

# 个人信息输入
st.header("Personal Information")
col1, col2, col3 = st.columns(3)
with col1:
    gender = st.selectbox("Gender", ["Male", "Female", "Other"])
with col2:
    height = st.number_input("Height (cm)", min_value=50, max_value=250, value=170)
with col3:
    weight = st.number_input("Weight (kg)", min_value=20, max_value=300, value=70)

# 食物输入
st.header("Food Entry")
food_type = st.text_input("Food Type")
amount = st.number_input("Amount (g)", min_value=0.0)

# 添加食物记录按钮
if st.button("Add Food Entry") and food_type and amount:
    entry = {
        'amount': amount,
        'food_type': food_type,
        'gender': gender,
        'height': height,
        'weight': weight
    }
    st.session_state.food_entries.append(entry)
    st.rerun()

# 显示食物记录和分析
if st.session_state.food_entries:
    df = pd.DataFrame(st.session_state.food_entries)
    st.subheader("Food Entries")
    st.write(df)

    # 基础营养计算
    total_calories = df['amount'].sum() * 2  # 假设每克食物2卡路里
    total_protein = df['amount'].sum() * 0.1  # 假设每克食物0.1克蛋白
    bmr = 10 * weight + 6.25 * height - 5 * 30 + (5 if gender == "Male" else -161)  # Assuming age 30

    # 营养图表
    chart_data = pd.DataFrame({
        'Nutrient': ['Calories', 'Protein'],
        'Amount': [total_calories, total_protein]
    })
    fig = px.bar(chart_data, x='Nutrient', y='Amount', title='Nutritional Breakdown')
    st.plotly_chart(fig)

    # 营养分析
    st.subheader("Nutritional Analysis")
    st.write(f"Estimated BMR: {bmr:.0f} calories")

    # AI Suggestions Button
    if st.button("Generate AI Nutrition Report"):
        if amount and food_type and height and weight and gender:
            entry = {
                'amount': amount,
                'gender': gender,
                'height': height,
                'weight': weight,
                'food_type': food_type

            }
            st.session_state.user_entries.append(entry)


            try:
                completion = jamai.add_table_rows(
                    "action",
                    p.RowAddRequest(
                        table_id="meal_analysis",
                        data=[entry],
                        stream=False
                    )
                )
                # 处理并显示返回的数据
                if completion.rows:
                    output_row = completion.rows[0].columns
                    suggestion = output_row.get("suggestion")
                    summary = output_row.get("summary")
                    nutrition_analysis = output_row.get("nutrition_analysis")



                    st.markdown(
                        f"""
                                   <div class="generated-output">
                                    <h4>📊 Recommended Diet:</h4> <p>{suggestion.text if suggestion else 'N/A'}</p>  
                                    <h4>🍽️ Nutrition Analysis:</h4> <p>{nutrition_analysis.text if nutrition_analysis else 'N/A'}</p>
                                    <h4>📝 Summary:</h4> <p>{summary.text if summary else 'N/A'}</p>
                                   </div>
                                   """,
                        unsafe_allow_html=True
                    )
                    with st.container():
                        st.subheader("📥 Download Final Report")

                        doc = Document()

                        # nutrition report
                        doc.add_heading("Nutrition Report", level=1)
                        doc.add_paragraph(nutrition_analysis.text if nutrition_analysis else 'N/A')

                        # Suggestion part
                        doc.add_heading("Recommended Diet", level=2)
                        doc.add_paragraph(suggestion.text if suggestion else 'N/A')

                        # Summary
                        doc.add_heading("Conclusion", level=2)  # Changed from add_paragraph to add_heading
                        doc.add_paragraph(summary.text if summary else 'N/A')

                        buffer = BytesIO()
                        doc.save(buffer)
                        buffer.seek(0)

                        st.download_button(
                            label="📄 Download Final Report as .docx",
                            data=buffer,
                            file_name=generate_random_filename(),
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )

            except Exception as e:
                st.error(f"Error processing the nutrition report: {e}")
        else:
            st.write("Add food entries to get personalized nutrition analysis!")